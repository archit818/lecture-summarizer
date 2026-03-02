"""
Post-session summarizer.
Combines transcripts, sends to Ollama for summary, generates .docx document.
"""

import logging
import requests
from datetime import datetime
from pathlib import Path

from docx import Document

from . import config, database

logger = logging.getLogger(__name__)


def _build_prompt(full_transcript: str) -> str:
    """Build a structured prompt for the LLM."""
    return f"""You are an AI assistant that summarizes lecture and meeting transcripts.

Given the following transcript, produce a structured summary with these sections:

1. **Summary** — A concise overview of what was discussed (3–5 sentences).
2. **Key Points** — A bullet list of the most important points, decisions, or takeaways.
3. **Important Announcements** — Any deadlines, action items, or announcements mentioned.

If a section has no relevant content, write "None identified."

--- TRANSCRIPT START ---
{full_transcript}
--- TRANSCRIPT END ---

Respond in plain text with clear section headers."""


def generate_summary(full_transcript: str) -> str:
    """Send transcript to Ollama and get a structured summary."""
    if not full_transcript.strip():
        return "No transcript content available for summarization."

    prompt = _build_prompt(full_transcript)

    try:
        resp = requests.post(
            f"{config.OLLAMA_URL}/api/generate",
            json={
                "model": config.OLLAMA_MODEL,
                "prompt": prompt,
                "stream": False,
            },
            timeout=300,  # LLM can be slow
        )
        resp.raise_for_status()
        result = resp.json()
        summary = result.get("response", "").strip()
        if summary:
            logger.info("Summary generated successfully.")
            return summary
        else:
            logger.warning("Ollama returned empty response.")
            return "Summary generation returned empty result."

    except requests.ConnectionError:
        logger.error("Cannot connect to Ollama. Is it running?")
        return "Summary unavailable — Ollama not reachable."
    except Exception as e:
        logger.error(f"Summary generation failed: {e}", exc_info=True)
        return f"Summary generation failed: {e}"


def generate_document(session_id: int) -> str:
    """
    Generate a .docx document for the completed session.
    Returns the file path.
    """
    # Gather data
    session = database.get_active_session()
    if session is None:
        # Session might already be ended — query by ID
        from . import database as db
        conn = db._get_connection()
        try:
            row = conn.execute("SELECT * FROM sessions WHERE id = ?", (session_id,)).fetchone()
            session = dict(row) if row else {}
        finally:
            conn.close()

    transcripts = database.get_session_transcripts(session_id)
    qr_codes = database.get_session_qr_codes(session_id)

    # Build full transcript text
    full_transcript = "\n".join(t["text"] for t in transcripts)

    # Generate summary
    summary_text = generate_summary(full_transcript)

    # Calculate duration
    start_time = session.get("start_time", "Unknown")
    end_time = session.get("end_time", "Unknown")
    try:
        start_dt = datetime.fromisoformat(start_time)
        end_dt = datetime.fromisoformat(end_time)
        duration = str(end_dt - start_dt).split(".")[0]  # remove microseconds
    except Exception:
        duration = "Unknown"

    # Build docx
    doc = Document()
    doc.add_heading("Lecture / Meeting Summary", level=0)

    doc.add_heading("Session Details", level=1)
    doc.add_paragraph(f"Date: {start_time[:10] if len(start_time) >= 10 else start_time}")
    doc.add_paragraph(f"Start: {start_time}")
    doc.add_paragraph(f"End: {end_time}")
    doc.add_paragraph(f"Duration: {duration}")
    doc.add_paragraph(f"Session ID: {session_id}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary_text)

    if qr_codes:
        doc.add_heading("QR Code Links", level=1)
        # Deduplicate for display
        seen_urls = set()
        for qr in qr_codes:
            url = qr["url"]
            if url not in seen_urls:
                doc.add_paragraph(f"• {url}", style="List Bullet")
                seen_urls.add(url)

    doc.add_heading("Full Transcript", level=1)
    if full_transcript.strip():
        # Add in chunks to avoid huge paragraphs
        for t in transcripts:
            ts = t["timestamp"]
            doc.add_paragraph(f"[{ts}] {t['text']}")
    else:
        doc.add_paragraph("No transcript recorded.")

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"session_{session_id}_{timestamp}.docx"
    filepath = config.OUTPUT_DIR / filename
    doc.save(str(filepath))
    logger.info(f"Document generated: {filepath}")
    return str(filepath)
